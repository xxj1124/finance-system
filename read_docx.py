from docx import Document

def read_docx_content(file_path):
    try:
        doc = Document(file_path)
        content = []
        
        # 读取段落
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        
        # 读取表格
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                if any(row_text):
                    content.append('\t'.join(row_text))
        
        return '\n'.join(content)
    except Exception as e:
        print(f"读取文档失败: {e}")
        return None

if __name__ == "__main__":
    file_path = "1602230103  杨深.docx"
    content = read_docx_content(file_path)
    if content:
        with open("document_content.txt", "w", encoding="utf-8") as f:
            f.write(content)
        print("文档内容已保存到document_content.txt")
